#!/usr/bin/env python3
"""
测试Python解析服务
"""

import sys
import os
import tempfile
import logging

# 添加当前目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 导入解析器
from service.pdf_parser import PDFParser
from service.docx_parser import DOCXParser

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def create_test_pdf():
    """创建一个简单的测试PDF文件"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        # 创建临时PDF文件
        temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        temp_pdf.close()
        
        # 创建PDF内容
        c = canvas.Canvas(temp_pdf.name, pagesize=letter)
        c.drawString(100, 750, "这是一个测试PDF文档")
        c.drawString(100, 730, "用于测试PDF解析功能")
        c.drawString(100, 710, "包含多行文本内容")
        c.save()
        
        logger.info(f"创建测试PDF文件: {temp_pdf.name}")
        return temp_pdf.name
        
    except ImportError:
        logger.warning("reportlab库未安装，跳过PDF测试")
        return None

def create_test_docx():
    """创建一个简单的测试DOCX文件"""
    try:
        from docx import Document
        
        # 创建临时DOCX文件
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_docx.close()
        
        # 创建DOCX内容
        doc = Document()
        doc.add_heading('测试DOCX文档', 0)
        doc.add_paragraph('这是一个测试DOCX文档')
        doc.add_paragraph('用于测试DOCX解析功能')
        doc.add_paragraph('包含多行文本内容')
        
        # 添加一个表格
        table = doc.add_table(rows=3, cols=2)
        for i in range(3):
            row = table.rows[i]
            row.cells[0].text = f'单元格 {i+1}-1'
            row.cells[1].text = f'单元格 {i+1}-2'
        
        doc.save(temp_docx.name)
        
        logger.info(f"创建测试DOCX文件: {temp_docx.name}")
        return temp_docx.name
        
    except ImportError:
        logger.warning("python-docx库未安装，跳过DOCX测试")
        return None

def test_pdf_parser():
    """测试PDF解析器"""
    logger.info("开始测试PDF解析器")
    
    pdf_file = create_test_pdf()
    if not pdf_file:
        logger.warning("无法创建测试PDF文件，跳过PDF测试")
        return
    
    try:
        parser = PDFParser()
        content, metadata = parser.parse(pdf_file)
        
        logger.info(f"PDF解析成功")
        logger.info(f"内容长度: {len(content)}")
        logger.info(f"元数据: {metadata}")
        
        # 清理临时文件
        os.unlink(pdf_file)
        
    except Exception as e:
        logger.error(f"PDF解析测试失败: {e}")
        if os.path.exists(pdf_file):
            os.unlink(pdf_file)

def test_docx_parser():
    """测试DOCX解析器"""
    logger.info("开始测试DOCX解析器")
    
    docx_file = create_test_docx()
    if not docx_file:
        logger.warning("无法创建测试DOCX文件，跳过DOCX测试")
        return
    
    try:
        parser = DOCXParser()
        content, metadata = parser.parse(docx_file)
        
        logger.info(f"DOCX解析成功")
        logger.info(f"内容长度: {len(content)}")
        logger.info(f"元数据: {metadata}")
        
        # 清理临时文件
        os.unlink(docx_file)
        
    except Exception as e:
        logger.error(f"DOCX解析测试失败: {e}")
        if os.path.exists(docx_file):
            os.unlink(docx_file)

def main():
    """主函数"""
    logger.info("开始测试Python解析服务")
    
    # 测试PDF解析器
    test_pdf_parser()
    
    # 测试DOCX解析器
    test_docx_parser()
    
    logger.info("测试完成")

if __name__ == "__main__":
    main()