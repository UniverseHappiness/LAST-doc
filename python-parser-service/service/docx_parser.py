import os
import logging
from typing import Dict, Any, Tuple
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

logger = logging.getLogger(__name__)

class DOCXParser:
    """DOCX文档解析器"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def parse(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        解析DOCX文档
        
        Args:
            file_path: DOCX文件路径
            
        Returns:
            Tuple[str, Dict[str, Any]]: (文本内容, 元数据)
        """
        try:
            self.logger.info(f"开始解析DOCX文档: {file_path}")
            
            # 检查文件是否存在
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"DOCX文件不存在: {file_path}")
            
            # 获取文件大小
            file_size = os.path.getsize(file_path)
            self.logger.info(f"DOCX文件大小: {file_size} 字节")
            
            # 加载DOCX文档
            doc = Document(file_path)
            
            # 提取文本内容
            content_text = self._extract_text_content(doc)
            
            # 提取元数据
            metadata = self._extract_metadata(doc, file_path, file_size)
            
            self.logger.info(f"DOCX文档解析完成，内容长度: {len(content_text)}")
            
            return content_text, metadata
            
        except Exception as e:
            self.logger.error(f"解析DOCX文档时发生错误: {e}")
            raise
    
    def _extract_text_content(self, doc: Document) -> str:
        """提取文档文本内容"""
        content_parts = []
        
        # 提取段落文本
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                content_parts.append(paragraph.text)
                self.logger.debug(f"已处理第 {i + 1} 个段落，文本长度: {len(paragraph.text)}")
        
        # 提取表格文本
        for i, table in enumerate(doc.tables):
            table_text = self._extract_table_text(table)
            if table_text:
                content_parts.append(f"[表格 {i + 1}]\n{table_text}")
                self.logger.debug(f"已处理第 {i + 1} 个表格")
        
        # 提取页眉页脚文本
        for section in doc.sections:
            # 页眉
            for i, header in enumerate(section.header.paragraphs):
                if header.text.strip():
                    content_parts.append(f"[页眉 {i + 1}] {header.text}")
            
            # 页脚
            for i, footer in enumerate(section.footer.paragraphs):
                if footer.text.strip():
                    content_parts.append(f"[页脚 {i + 1}] {footer.text}")
        
        return "\n".join(content_parts)
    
    def _extract_table_text(self, table) -> str:
        """提取表格文本内容"""
        table_rows = []
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_cells.append(cell_text)
            if row_cells:
                table_rows.append(" | ".join(row_cells))
        
        return "\n".join(table_rows)
    
    def _extract_metadata(self, doc: Document, file_path: str, file_size: int) -> Dict[str, Any]:
        """提取文档元数据"""
        metadata = {
            'file_size': str(file_size),
            'paragraph_count': str(len(doc.paragraphs)),
            'table_count': str(len(doc.tables)),
            'section_count': str(len(doc.sections)),
            'content_length': str(len(self._extract_text_content(doc))),
            'parser': 'python-docx'
        }
        
        # 提取核心属性
        core_props = doc.core_properties
        if core_props:
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'language': core_props.language or '',
                'category': core_props.category or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'last_modified_by': core_props.last_modified_by or '',
                'revision': str(core_props.revision) if core_props.revision else '',
                'version': core_props.version or ''
            })
        
        # 统计图片数量
        image_count = 0
        for rel in doc.part.rels.values():
            if rel.reltype == RT.IMAGE:
                image_count += 1
        metadata['image_count'] = str(image_count)
        
        # 统计超链接数量
        hyperlink_count = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if hasattr(run, 'hyperlink') and run.hyperlink:
                    hyperlink_count += 1
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if hasattr(run, 'hyperlink') and run.hyperlink:
                                hyperlink_count += 1
        
        metadata['hyperlink_count'] = str(hyperlink_count)
        
        return metadata