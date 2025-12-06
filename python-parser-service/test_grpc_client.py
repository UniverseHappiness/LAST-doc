#!/usr/bin/env python3
"""
测试gRPC客户端
"""

import sys
import os
import logging
import grpc
import tempfile

# 添加当前目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 导入生成的gRPC代码
try:
    import document_parser_pb2 as pb2
    import document_parser_pb2_grpc as pb2_grpc
except ImportError:
    print("错误：请先生成gRPC代码")
    sys.exit(1)

# 导入解析器
from service.docx_parser import DOCXParser

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def create_test_docx():
    """创建一个简单的测试DOCX文件"""
    try:
        from docx import Document
        
        # 创建临时DOCX文件
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_docx.close()
        
        # 创建DOCX内容
        doc = Document()
        doc.add_heading('测试DOCX文档', 0)
        doc.add_paragraph('这是一个测试DOCX文档')
        doc.add_paragraph('用于测试gRPC通信')
        doc.add_paragraph('包含多行文本内容')
        
        # 添加一个表格
        table = doc.add_table(rows=3, cols=2)
        for i in range(3):
            row = table.rows[i]
            row.cells[0].text = f'单元格 {i+1}-1'
            row.cells[1].text = f'单元格 {i+1}-2'
        
        doc.save(temp_docx.name)
        
        logger.info(f"创建测试DOCX文件: {temp_docx.name}")
        return temp_docx.name
        
    except ImportError:
        logger.warning("python-docx库未安装，跳过DOCX测试")
        return None

def test_grpc_connection():
    """测试gRPC连接"""
    logger.info("开始测试gRPC连接")
    
    try:
        # 创建gRPC通道
        channel = grpc.insecure_channel('127.0.0.1:50051')
        
        # 创建客户端存根
        stub = pb2_grpc.DocumentParserServiceStub(channel)
        
        # 测试健康检查
        logger.info("测试健康检查...")
        health_request = pb2.HealthCheckRequest(service="document_parser")
        health_response = stub.HealthCheck(health_request)
        
        logger.info(f"健康检查结果: {health_response.healthy}")
        logger.info(f"健康检查消息: {health_response.message}")
        logger.info(f"服务版本: {health_response.version}")
        
        # 测试DOCX解析
        logger.info("测试DOCX解析...")
        docx_file = create_test_docx()
        if docx_file:
            docx_request = pb2.ParseDOCXRequest(file_path=docx_file)
            docx_response = stub.ParseDOCX(docx_request)
            
            logger.info(f"DOCX解析成功: {docx_response.success}")
            if docx_response.success:
                logger.info(f"内容长度: {len(docx_response.content)}")
                logger.info(f"元数据: {dict(docx_response.metadata)}")
            else:
                logger.error(f"DOCX解析失败: {docx_response.error_message}")
            
            # 清理临时文件
            os.unlink(docx_file)
        
        # 关闭通道
        channel.close()
        
        logger.info("gRPC连接测试完成")
        
    except grpc.RpcError as e:
        logger.error(f"gRPC连接失败: {e}")
        logger.error("这可能是因为Python gRPC服务器没有启动")
        logger.info("请先运行: python -m service.server")

def main():
    """主函数"""
    logger.info("开始测试gRPC客户端")
    
    # 测试gRPC连接
    test_grpc_connection()
    
    logger.info("测试完成")

if __name__ == "__main__":
    main()